"""Document parsing for RAG ingestion.

Supports multiple document formats:
- Plain text (.txt)
- Markdown (.md)
- HTML (.html, .htm)
- PDF (.pdf)
- Word documents (.docx)
- JSON (.json)
"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parses documents into plain text for chunking."""

    SUPPORTED_EXTENSIONS = {
        ".txt",
        ".md",
        ".markdown",
        ".html",
        ".htm",
        ".pdf",
        ".docx",
        ".json",
    }

    def parse(self, content: str | bytes, file_path: str | None = None) -> dict[str, Any]:
        """Parse document content.

        Args:
            content: Document content (str or bytes).
            file_path: Optional file path to determine format.

        Returns:
            Dict with 'text', 'metadata', and 'chunks'.
        """
        # Determine format from file extension or content
        if file_path:
            ext = Path(file_path).suffix.lower()
            mime_type = self._get_mime_type(ext)
        else:
            mime_type = self._detect_mime_type(content)
            ext = ""

        logger.debug(f"Parsing document", extra={"mime_type": mime_type, "path": file_path})

        # Parse based on mime type
        if mime_type == "text/markdown" or ext in (".md", ".markdown"):
            return self._parse_markdown(content, file_path)
        elif mime_type == "text/html" or ext in (".html", ".htm"):
            return self._parse_html(content, file_path)
        elif mime_type == "application/pdf" or ext == ".pdf":
            return self._parse_pdf(content, file_path)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or ext == ".docx":
            return self._parse_docx(content, file_path)
        elif mime_type == "application/json" or ext == ".json":
            return self._parse_json(content, file_path)
        else:
            # Default to plain text
            return self._parse_text(content, file_path)

    def _get_mime_type(self, ext: str) -> str:
        """Get MIME type from file extension."""
        mime_types = {
            ".txt": "text/plain",
            ".md": "text/markdown",
            ".markdown": "text/markdown",
            ".html": "text/html",
            ".htm": "text/html",
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".json": "application/json",
        }
        return mime_types.get(ext, "text/plain")

    def _detect_mime_type(self, content: str | bytes) -> str:
        """Detect MIME type from content."""
        if isinstance(content, bytes):
            # Check for PDF magic number
            if content.startswith(b"%PDF"):
                return "application/pdf"
            # Try to decode as text
            try:
                content = content.decode("utf-8")
            except UnicodeDecodeError:
                return "application/octet-stream"

        # Check for HTML tags
        if re.search(r"<\s*(html|body|div|p|span)", content, re.IGNORECASE):
            return "text/html"

        # Check for JSON
        if content.strip().startswith(("{", "[")):
            try:
                json.loads(content)
                return "application/json"
            except json.JSONDecodeError:
                pass

        # Default to plain text
        return "text/plain"

    def _parse_text(self, content: str | bytes, file_path: str | None = None) -> dict[str, Any]:
        """Parse plain text."""
        if isinstance(content, bytes):
            text = content.decode("utf-8", errors="ignore")
        else:
            text = content

        return {
            "text": text,
            "metadata": {
                "source": file_path,
                "format": "text/plain",
            },
        }

    def _parse_markdown(self, content: str | bytes, file_path: str | None = None) -> dict[str, Any]:
        """Parse Markdown document."""
        if isinstance(content, bytes):
            text = content.decode("utf-8", errors="ignore")
        else:
            text = content

        # Remove frontmatter (YAML between --- markers)
        text = re.sub(r"^---\s*\n.*?\n---\s*\n", "", text, flags=re.DOTALL)

        # Remove markdown formatting for plain text extraction
        # Keep headers and structure markers
        clean_text = text

        # Extract title from first heading
        title_match = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
        title = title_match.group(1) if title_match else None

        return {
            "text": clean_text,
            "metadata": {
                "source": file_path,
                "format": "text/markdown",
                "title": title,
            },
        }

    def _parse_html(self, content: str | bytes, file_path: str | None = None) -> dict[str, Any]:
        """Parse HTML document."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.warning("beautifulsoup4 not installed, using basic HTML parsing")
            # Fallback: use regex to remove tags
            if isinstance(content, bytes):
                text = content.decode("utf-8", errors="ignore")
            else:
                text = content
            text = re.sub(r"<[^>]+>", " ", text)
            text = re.sub(r"\s+", " ", text).strip()
            return {
                "text": text,
                "metadata": {
                    "source": file_path,
                    "format": "text/html",
                },
            }

        if isinstance(content, bytes):
            soup = BeautifulSoup(content, "html.parser", from_encoding="utf-8")
        else:
            soup = BeautifulSoup(content, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer"]):
            script.decompose()

        # Extract title
        title = None
        title_tag = soup.find("title")
        if title_tag:
            title = title_tag.get_text(strip=True)
        else:
            h1 = soup.find("h1")
            if h1:
                title = h1.get_text(strip=True)

        # Get text
        text = soup.get_text(separator="\n", strip=True)

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)

        return {
            "text": text,
            "metadata": {
                "source": file_path,
                "format": "text/html",
                "title": title,
            },
        }

    def _parse_pdf(self, content: bytes, file_path: str | None = None) -> dict[str, Any]:
        """Parse PDF document."""
        try:
            import pdfplumber
        except ImportError:
            logger.warning("pdfplumber not installed, returning empty text")
            return {
                "text": "",
                "metadata": {
                    "source": file_path,
                    "format": "application/pdf",
                    "error": "pdfplumber not installed",
                },
            }

        try:
            text_parts = []
            metadata = {"source": file_path, "format": "application/pdf"}

            with pdfplumber.open(BytesIO(content)) as pdf:
                # Extract metadata
                if pdf.metadata:
                    metadata["title"] = pdf.metadata.get("Title")
                    metadata["author"] = pdf.metadata.get("Author")

                # Extract text from each page
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"\n--- Page {i + 1} ---\n")
                        text_parts.append(page_text)

            return {
                "text": "\n".join(text_parts),
                "metadata": metadata,
            }

        except Exception as e:
            logger.error(f"Failed to parse PDF", extra={"error": str(e)})
            return {
                "text": "",
                "metadata": {
                    "source": file_path,
                    "format": "application/pdf",
                    "error": str(e),
                },
            }

    def _parse_docx(self, content: bytes, file_path: str | None = None) -> dict[str, Any]:
        """Parse Word document."""
        try:
            import docx
        except ImportError:
            logger.warning("python-docx not installed, returning empty text")
            return {
                "text": "",
                "metadata": {
                    "source": file_path,
                    "format": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "error": "python-docx not installed",
                },
            }

        try:
            from io import BytesIO

            doc = docx.Document(BytesIO(content))

            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            # Get document properties
            metadata = {
                "source": file_path,
                "format": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            }

            if doc.core_properties:
                metadata["title"] = doc.core_properties.title
                metadata["author"] = doc.core_properties.author

            return {
                "text": "\n\n".join(text_parts),
                "metadata": metadata,
            }

        except Exception as e:
            logger.error(f"Failed to parse DOCX", extra={"error": str(e)})
            return {
                "text": "",
                "metadata": {
                    "source": file_path,
                    "format": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "error": str(e),
                },
            }

    def _parse_json(self, content: str | bytes, file_path: str | None = None) -> dict[str, Any]:
        """Parse JSON document."""
        if isinstance(content, bytes):
            text = content.decode("utf-8", errors="ignore")
        else:
            text = content

        try:
            data = json.loads(text)
            # Convert to formatted text
            formatted = json.dumps(data, indent=2, ensure_ascii=False)

            return {
                "text": formatted,
                "metadata": {
                    "source": file_path,
                    "format": "application/json",
                },
            }
        except json.JSONDecodeError as e:
            logger.warning(f"Invalid JSON", extra={"error": str(e)})
            return {
                "text": text,
                "metadata": {
                    "source": file_path,
                    "format": "application/json",
                    "error": str(e),
                },
            }


def parse_file(file_path: str | Path) -> dict[str, Any]:
    """Parse a file from disk.

    Args:
        file_path: Path to the file.

    Returns:
        Parsed document data.
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    # Read file
    if file_path.suffix.lower() in (".pdf", ".docx"):
        mode = "rb"
    else:
        mode = "r"

    with open(file_path, mode, encoding="utf-8" if mode == "r" else None) as f:
        content = f.read()

    # Parse
    parser = DocumentParser()
    return parser.parse(content, str(file_path))


def parse_directory(
    directory: str | Path,
    extensions: set[str] | None = None,
) -> list[dict[str, Any]]:
    """Parse all supported files in a directory.

    Args:
        directory: Directory path.
        extensions: Set of extensions to include (default: all supported).

    Returns:
        List of parsed documents.
    """
    directory = Path(directory)

    if extensions is None:
        extensions = DocumentParser.SUPPORTED_EXTENSIONS

    results = []

    for file_path in directory.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in extensions:
            try:
                doc = parse_file(file_path)
                results.append(doc)
                logger.debug(f"Parsed {file_path}")
            except Exception as e:
                logger.error(f"Failed to parse {file_path}", extra={"error": str(e)})

    return results


# Import BytesIO for PDF parsing
from io import BytesIO
